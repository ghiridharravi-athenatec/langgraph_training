import os
import fitz
import pandas as pd

from typing import List
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_mongodb import MongoDBAtlasVectorSearch
from langchain_community.embeddings import HuggingFaceEmbeddings
import torch
from app.utils.mongo import DOCUMENT_CHUNKS_COLLECTION, get_mongo_client, create_vector_search_index
from app.utils.table_chunking import chunk_table_rows
import io
from PIL import Image
from paddleocr import PaddleOCR
from app.core.ingest_guardrails import scan_ingested_pii
from app.core.logger import get_logger

logger = get_logger(__name__)

ocr = PaddleOCR(
    use_doc_orientation_classify=False,
    use_doc_unwarping=False,
    use_textline_orientation=False,
    lang="en",
    # oneDNN's PIR executor path crashes on this CPU with
    # "ConvertPirAttribute2RuntimeAttribute not support [pir::ArrayAttribute<pir::DoubleAttribute>]"
    # on every image; the plain (non-mkldnn) run mode avoids that op path.
    enable_mkldnn=False,
)

device = "cuda" if torch.cuda.is_available() else "mps" if torch.backends.mps.is_available() else "cpu"
embedding_model = HuggingFaceEmbeddings(
    model_name="BAAI/bge-m3",
    model_kwargs={"device": device},
    encode_kwargs={"normalize_embeddings": True},
)

def extract_images(pdf: fitz.Document, pdf_path: str):

    docs = []

    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]

    image_dir = os.path.join(
        "app",
        "extracted_images",
        pdf_name,
    )

    os.makedirs(image_dir, exist_ok=True)

    for page_no, page in enumerate(pdf, start=1):

        images = page.get_images(full=True)

        if not images:
            continue

        page_text = page.get_text("text")

        for img_index, img in enumerate(images):

            xref = img[0]

            base_image = pdf.extract_image(xref)

            image_bytes = base_image["image"]

            ext = base_image["ext"]

            image_name = f"page_{page_no}_img_{img_index}.{ext}"

            image_path = os.path.join(image_dir, image_name)

            with open(image_path, "wb") as f:
                f.write(image_bytes)

            # OCR
            ocr_text = ""

            try:

                result = ocr.predict(image_path)

                if result:

                    texts = []

                    for res in result:

                        if "rec_texts" in res:

                            texts.extend(res["rec_texts"])

                    ocr_text = "\n".join(texts)

            except Exception as e:
                logger.exception("OCR failed on '%s': %s", image_path, e)

            docs.append(
                Document(
                    page_content=f"""
                                Page Text:
                                {page_text}

                                Image OCR:
                                {ocr_text}
                                """,
                    metadata={
                        "source": os.path.basename(pdf_path),
                        "page": page_no,
                        "sheet_name": "",
                        "content_type": "pdf_image",
                        "image_path": image_path,
                    },
                )
            )

    return docs


def load_pdf(pdf_path: str) -> List[Document]:
    docs = []
    pdf = fitz.open(pdf_path)

    image_docs = extract_images(pdf, pdf_path)

    for page_no, page in enumerate(pdf, start=1):
        text = page.get_text("text")

        if text.strip():
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": os.path.basename(pdf_path),
                        "page": page_no,
                        "sheet_name": "",
                        "content_type": "pdf_text",
                    },
                )
            )

        try:
            tables = page.find_tables()
            for table in tables:
                table_data = table.extract()
                if not table_data:
                    continue

                formatted_rows = [
                    " | ".join([str(cell) if cell else "" for cell in row])
                    for row in table_data
                ]
                if not any(row.strip() for row in formatted_rows):
                    continue

                # First row is the table's header - repeated into every chunk below
                # (chunk_table_rows) instead of only surviving in the first one.
                header_lines, row_lines = formatted_rows[:1], formatted_rows[1:]
                docs.extend(
                    chunk_table_rows(
                        header_lines, row_lines,
                        metadata={
                            "source": os.path.basename(pdf_path),
                            "page": page_no,
                            "sheet_name": "",
                            "content_type": "pdf_table",
                        },
                    )
                )
        except Exception as e:
            logger.exception("Failed to extract tables from page %s of '%s': %s", page_no, pdf_path, e)

    docs.extend(image_docs)

    return docs


def load_xlsx(xlsx_path: str) -> List[Document]:
    docs = []
    sheets = pd.read_excel(xlsx_path, sheet_name=None)
    source = os.path.basename(xlsx_path)

    for sheet_name, df in sheets.items():
        if df.empty:
            continue

        # to_markdown()'s first two lines are the column-name row and the markdown
        # separator row (|:---|---:|) - both are the "header" here and get repeated
        # into every chunk below (chunk_table_rows), not just the first one.
        lines = df.to_markdown(index=False).split("\n")
        header_lines, row_lines = (lines[:2], lines[2:]) if len(lines) >= 2 else (lines, [])

        docs.extend(
            chunk_table_rows(
                header_lines, row_lines,
                metadata={
                    "source": source,
                    "page": 0,
                    "sheet_name": sheet_name,
                    "content_type": "xlsx_table",
                },
            )
        )

    return docs


def load_docx(docx_path: str) -> List[Document]:
    docs = []
    doc = DocxDocument(docx_path)
    source = os.path.basename(docx_path)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        docs.append(
            Document(
                page_content="\n\n".join(paragraphs),
                metadata={"source": source, "page": 0, "sheet_name": "", "content_type": "docx_text"},
            )
        )

    for table in doc.tables:
        rows = ["\t".join(cell.text for cell in row.cells) for row in table.rows]
        if not rows or not any(row.strip() for row in rows):
            continue

        # First row is the table's header - repeated into every chunk below
        # (chunk_table_rows) instead of only surviving in the first one.
        header_lines, row_lines = rows[:1], rows[1:]
        docs.extend(
            chunk_table_rows(
                header_lines, row_lines,
                metadata={"source": source, "page": 0, "sheet_name": "", "content_type": "docx_table"},
            )
        )

    return docs


def load_txt(txt_path: str) -> List[Document]:
    with open(txt_path, "r", encoding="utf-8") as f:
        content = f.read()

    if not content.strip():
        return []

    return [
        Document(
            page_content=content,
            metadata={"source": os.path.basename(txt_path), "page": 0, "sheet_name": "", "content_type": "text"},
        )
    ]


_LOADERS = {
    ".pdf": load_pdf,
    ".xlsx": load_xlsx,
    ".docx": load_docx,
    ".txt": load_txt,
}

# Table-derived documents (from load_xlsx/load_pdf/load_docx) are already
# chunked row-batch-wise by chunk_table_rows, with the header repeated into
# every batch - running the generic character-count splitter over them again
# would cut them again with no table awareness, silently reintroducing the
# same header-loss problem chunk_table_rows exists to avoid.
_TABLE_CONTENT_TYPES = {"pdf_table", "xlsx_table", "docx_table"}


def ingest_files(file_paths: List[str], user_id: str, pii_entities: List[str] = None):
    try:
        logger.info("Starting ingestion of %d file(s) for user %s", len(file_paths), user_id)
        all_docs = []

        for path in file_paths:
            extension = os.path.splitext(path)[1].lower()
            loader = _LOADERS.get(extension)
            if loader is not None:
                all_docs.extend(loader(path))

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=700,
            chunk_overlap=200,
            separators=[
                "\n\n",
                "\n",
                ". ",
                " ",
                ""
            ]
        )

        table_docs = [d for d in all_docs if d.metadata.get("content_type") in _TABLE_CONTENT_TYPES]
        other_docs = [d for d in all_docs if d.metadata.get("content_type") not in _TABLE_CONTENT_TYPES]
        chunks = splitter.split_documents(other_docs) + table_docs

        # Stamped onto every chunk so retrieval can filter to this uploader's own
        # documents only (see retrieve.py's pre_filter / per-user BM25 cache) - nobody,
        # including admins, retrieves chunks another user uploaded.
        for chunk in chunks:
            chunk.metadata["user_id"] = user_id

        pii_event = scan_ingested_pii(chunks, entities=pii_entities)

        client = get_mongo_client()
        collection = client["rag_database"][DOCUMENT_CHUNKS_COLLECTION]

        vectorstore = MongoDBAtlasVectorSearch.from_documents(
            documents=chunks,
            embedding=embedding_model,
            collection=collection,
            index_name="default",
        )

        create_vector_search_index(DOCUMENT_CHUNKS_COLLECTION)

        logger.info("Ingestion completed. Total chunks: %d", len(chunks))

        # scan_ingested_pii already redacted chunk.page_content in place before this
        # point, so the joined text below is already PII-masked - used as-is for the
        # "extracted content" preview stored on the document record.
        extracted_text = "\n\n".join(chunk.page_content for chunk in chunks)

        return {
            "passed": True,
            "message": f"Document ingested successfully. Total chunks: {len(chunks)}",
            "pii_event": pii_event,
            "chunk_count": len(chunks),
            "extracted_text": extracted_text,
        }

    except Exception as e:
        logger.exception("Error during ingestion: %s", e)
        return {"passed": False, "error": str(e)}